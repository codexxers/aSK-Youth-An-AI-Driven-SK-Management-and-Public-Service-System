from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import sys
import io
import sqlite3
import json

# ---------------------------------------------------------------------------
# Python AI Layer — Phase 6 Upgrade
# Provides NLP services called by the Node.js backend via HTTP.
# ALL models run on CPU (device=-1) — GPU VRAM is reserved for Qwen in Node.js.
# ---------------------------------------------------------------------------

app = FastAPI(title="aSK//YOUTH Python AI Layer", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------------------------------------------------------------------
# Module-level model holders — loaded ONCE at startup, reused per request
# ---------------------------------------------------------------------------
_intent_classifier = None
_summarizer_model = None
_summarizer_tokenizer = None
_embedding_model = None

@app.on_event("startup")
def load_models():
    """Load all NLP models at startup. All on CPU to avoid VRAM contention."""
    global _intent_classifier, _summarizer_model, _summarizer_tokenizer, _embedding_model

    # Feature 1: Intent Classification (facebook/bart-large-mnli)
    try:
        from transformers import pipeline as hf_pipeline
        print("[Python AI] Loading intent classifier (facebook/bart-large-mnli) on CPU...")
        _intent_classifier = hf_pipeline(
            "zero-shot-classification",
            model="facebook/bart-large-mnli",
            device=-1  # CPU only
        )
        print("[Python AI] Intent classifier ready.")
    except Exception as e:
        print(f"[Python AI] WARNING: Intent classifier failed to load: {e}")

    # Feature 3: Summarization (sshleifer/distilbart-cnn-12-6)
    # Using model + tokenizer directly since transformers v5 removed the
    # "summarization" pipeline task name.
    try:
        from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
        _summarizer_model_name = "sshleifer/distilbart-cnn-12-6"
        print(f"[Python AI] Loading summarizer ({_summarizer_model_name}) on CPU...")
        _summarizer_tokenizer = AutoTokenizer.from_pretrained(_summarizer_model_name)
        _summarizer_model = AutoModelForSeq2SeqLM.from_pretrained(_summarizer_model_name)
        _summarizer_model.eval()
        print("[Python AI] Summarizer ready.")
    except Exception as e:
        print(f"[Python AI] WARNING: Summarizer failed to load: {e}")

    # Feature 6: Embedding model (all-MiniLM-L6-v2)
    try:
        from sentence_transformers import SentenceTransformer
        print("[Python AI] Loading embedding model (all-MiniLM-L6-v2) on CPU...")
        _embedding_model = SentenceTransformer("all-MiniLM-L6-v2", device="cpu")
        print("[Python AI] Embedding model ready.")
    except Exception as e:
        print(f"[Python AI] WARNING: Embedding model failed to load: {e}")

# ---------------------------------------------------------------------------
# Health Check
# ---------------------------------------------------------------------------
@app.get("/health")
def health():
    return {
        "status": "ok",
        "service": "python-ai-layer",
        "models": {
            "intent_classifier": _intent_classifier is not None,
            "summarizer": _summarizer_model is not None,
            "embedding_model": _embedding_model is not None,
        }
    }

# ---------------------------------------------------------------------------
# Feature 4: Language Detection
# ---------------------------------------------------------------------------
class LanguageRequest(BaseModel):
    text: str

@app.post("/detect-language")
def detect_language(req: LanguageRequest):
    if not req.text or not req.text.strip():
        return {"language": "en", "is_filipino": False}
    try:
        from langdetect import detect_langs
        results = detect_langs(req.text)
        top = results[0]
        lang_code = str(top.lang)
        is_filipino = lang_code in ("tl", "fil")
        return {
            "language": lang_code,
            "confidence": round(top.prob, 3),
            "is_filipino": is_filipino
        }
    except Exception as e:
        print(f"[Language Detection] Error: {e}")
        return {"language": "en", "confidence": 0.0, "is_filipino": False}

# ---------------------------------------------------------------------------
# Feature 1: Intent Classification
# ---------------------------------------------------------------------------
class IntentRequest(BaseModel):
    text: str

INTENT_LABELS = [
    "casual conversation",
    "professional or official request",
    "document or file analysis"
]
LABEL_TO_MODE = {
    "casual conversation": "A",
    "professional or official request": "B",
    "document or file analysis": "C"
}

@app.post("/classify-intent")
def classify_intent(req: IntentRequest):
    if not _intent_classifier:
        raise HTTPException(status_code=503, detail="Intent classifier not loaded")
    if not req.text or not req.text.strip():
        return {"intent_mode": "A", "confidence": 1.0}
    try:
        result = _intent_classifier(req.text, INTENT_LABELS)
        top_label = result["labels"][0]
        top_score = result["scores"][0]
        return {
            "intent_mode": LABEL_TO_MODE.get(top_label, "A"),
            "confidence": round(top_score, 3)
        }
    except Exception as e:
        print(f"[Intent Classification] Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ---------------------------------------------------------------------------
# Feature 6: Embedding Service
# ---------------------------------------------------------------------------
class EmbedRequest(BaseModel):
    texts: List[str]

@app.post("/embed")
def embed_texts(req: EmbedRequest):
    if not _embedding_model:
        raise HTTPException(status_code=503, detail="Embedding model not loaded")
    if not req.texts:
        return {"embeddings": []}
    try:
        embeddings = _embedding_model.encode(
            req.texts,
            normalize_embeddings=True,
            show_progress_bar=False
        )
        return {"embeddings": embeddings.tolist()}
    except Exception as e:
        print(f"[Embedding] Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ---------------------------------------------------------------------------
# Feature 3: Document Summarization
# ---------------------------------------------------------------------------
class SummarizeRequest(BaseModel):
    text: str
    max_length: Optional[int] = 200

@app.post("/summarize")
def summarize_text(req: SummarizeRequest):
    if not _summarizer_model or not _summarizer_tokenizer:
        raise HTTPException(status_code=503, detail="Summarizer not loaded")
    if not req.text or not req.text.strip():
        return {"summary": ""}
    try:
        import torch
        # Tokenize with truncation (distilbart-cnn-12-6 has a 1024-token input limit)
        inputs = _summarizer_tokenizer(
            req.text,
            return_tensors="pt",
            max_length=1024,
            truncation=True
        )
        with torch.no_grad():
            summary_ids = _summarizer_model.generate(
                inputs["input_ids"],
                max_length=req.max_length,
                min_length=30,
                do_sample=False,
                num_beams=4
            )
        summary = _summarizer_tokenizer.decode(summary_ids[0], skip_special_tokens=True)
        return {"summary": summary}
    except Exception as e:
        print(f"[Summarization] Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ---------------------------------------------------------------------------
# Feature 2: OCR for Scanned Documents  (Phase 2)
# ---------------------------------------------------------------------------
# Tesseract binary path — check local tools/ folder first, then system install
_PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
_TESSERACT_PATHS = [
    os.path.join(_PROJECT_ROOT, "tools", "Tesseract-OCR", "tesseract.exe"),  # local (portable)
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",                         # system install
]
_ocr_available = False
for _tess_path in _TESSERACT_PATHS:
    if os.path.exists(_tess_path):
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = _tess_path
        _ocr_available = True
        print(f"[Python AI] Tesseract OCR found at {_tess_path}")
        break
if not _ocr_available:
    print("[Python AI] WARNING: Tesseract not found - OCR endpoint disabled")

# Poppler binary path for pdf2image (PDF OCR) — check local tools/ folder
_poppler_path = None
_poppler_tools_dir = os.path.join(_PROJECT_ROOT, "tools", "poppler")
if os.path.isdir(_poppler_tools_dir):
    for _dirpath, _dirnames, _filenames in os.walk(_poppler_tools_dir):
        if "pdftoppm.exe" in _filenames:
            _poppler_path = _dirpath
            print(f"[Python AI] Poppler found at {_poppler_path}")
            break
if _poppler_path is None:
    print("[Python AI] Poppler not in tools/ - pdf2image will use system PATH")

IMAGE_MIME_TYPES = {"image/jpeg", "image/png", "image/tiff", "image/webp", "image/bmp"}

@app.post("/ocr")
async def ocr_extract(file: UploadFile = File(...)):
    if not _ocr_available:
        raise HTTPException(status_code=503, detail="Tesseract OCR is not installed on this server.")
    import pytesseract
    from PIL import Image

    content_type = file.content_type or ""
    file_bytes = await file.read()

    try:
        if content_type == "application/pdf":
            from pdf2image import convert_from_bytes
            pages = convert_from_bytes(file_bytes, dpi=300, poppler_path=_poppler_path)
            texts = [pytesseract.image_to_string(page) for page in pages]
            return {"extracted_text": "\n\n".join(texts), "pages": len(pages)}
        elif content_type in IMAGE_MIME_TYPES or content_type.startswith("image/"):
            img = Image.open(io.BytesIO(file_bytes))
            text = pytesseract.image_to_string(img)
            return {"extracted_text": text, "pages": 1}
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type for OCR: {content_type}")
    except HTTPException:
        raise
    except Exception as e:
        print(f"[OCR] Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ---------------------------------------------------------------------------
# Feature 7: Template-Based Document Generation  (Phase 2)
# ---------------------------------------------------------------------------
_TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")

class DocGenRequest(BaseModel):
    template_id: str
    data: Dict[str, Any]
    format: str = "docx"

VALID_TEMPLATES = {"resolution", "minutes", "certificate"}
VALID_FORMATS = {"docx", "pdf"}

@app.post("/generate-document")
def generate_document(req: DocGenRequest):
    if req.template_id not in VALID_TEMPLATES:
        raise HTTPException(status_code=400, detail=f"Invalid template_id. Must be one of: {', '.join(VALID_TEMPLATES)}")
    if req.format not in VALID_FORMATS:
        raise HTTPException(status_code=400, detail=f"Invalid format. Must be one of: {', '.join(VALID_FORMATS)}")

    from jinja2 import Environment, FileSystemLoader
    jinja_env = Environment(loader=FileSystemLoader(_TEMPLATES_DIR), autoescape=False)
    try:
        template = jinja_env.get_template(f"{req.template_id}.j2")
    except Exception:
        raise HTTPException(status_code=404, detail=f"Template '{req.template_id}.j2' not found in templates directory.")

    rendered = template.render(**req.data)

    if req.format == "docx":
        return _build_docx(rendered, req.template_id, req.data)
    else:
        return _build_pdf(rendered, req.template_id, req.data)

def _build_docx(rendered_text: str, template_id: str, data: dict):
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    import re as _docx_re

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    # Tight spacing globally — eliminates Word's default 8pt paragraph gaps
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # --- Page border (double-line, dark blue) ---
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'text')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'double')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '15')
        border.set(qn('w:color'), '1F3864')
        pgBorders.append(border)
    sectPr.append(pgBorders)

    # Letterhead — tight spacing (no gaps between lines)
    letterhead = [
        "Republic of the Philippines",
        "City of Marikina",
        "Barangay Concepcion Dos",
        "OFFICE OF THE SANGGUNIANG KABATAAN",
    ]
    for line in letterhead:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(11)
        if "SANGGUNIANG" in line:
            run.bold = True

    # Spacer between letterhead and title
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after = Pt(4)

    # Title from template_id
    title_map = {"resolution": "SK RESOLUTION", "minutes": "MEETING MINUTES", "certificate": "CERTIFICATE OF RECOGNITION"}
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run(title_map.get(template_id, template_id.upper()))
    title_run.bold = True
    title_run.font.size = Pt(14)

    # Spacer between title and body
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(4)
    sp2.paragraph_format.space_after = Pt(4)

    # Body — each line from rendered Jinja2 template
    # List items (- bullet or 1. numbered) get spacing; everything else is tight
    for line in rendered_text.split("\n"):
        p = doc.add_paragraph(line)
        p.style.font.size = Pt(12)
        stripped = line.strip()
        if stripped.startswith('- ') or _docx_re.match(r'^\d+\.\s', stripped):
            # List items — keep visible spacing between entries
            p.paragraph_format.space_after = Pt(4)
        elif stripped.isupper() and stripped.endswith(':'):
            # Section headers (ATTENDEES:, AGENDA:, etc.) — slight gap before
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        else:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{template_id}_{data.get('date', 'document')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

def _build_pdf(rendered_text: str, template_id: str, data: dict):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas as rl_canvas

    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=LETTER)
    width, height = LETTER
    _BORDER_MARGIN = 36  # 0.5 inch from page edge
    _BORDER_COLOR = (0.12, 0.22, 0.39)  # dark navy blue (#1F3864)

    def _draw_page_border(canvas):
        """Draw a double-line border on the current page."""
        canvas.setStrokeColor(_BORDER_COLOR)
        # Outer border
        canvas.setLineWidth(2)
        canvas.rect(_BORDER_MARGIN, _BORDER_MARGIN,
                    width - 2 * _BORDER_MARGIN, height - 2 * _BORDER_MARGIN)
        # Inner border (4pt inset)
        canvas.setLineWidth(0.75)
        canvas.rect(_BORDER_MARGIN + 4, _BORDER_MARGIN + 4,
                    width - 2 * (_BORDER_MARGIN + 4), height - 2 * (_BORDER_MARGIN + 4))

    _draw_page_border(c)

    # Letterhead
    letterhead = [
        "Republic of the Philippines",
        "City of Marikina",
        "Barangay Concepcion Dos",
        "OFFICE OF THE SANGGUNIANG KABATAAN",
    ]
    y = height - 72
    for line in letterhead:
        c.setFont("Helvetica-Bold" if "SANGGUNIANG" in line else "Helvetica", 11)
        c.drawCentredString(width / 2, y, line)
        y -= 16

    # Title
    y -= 20
    title_map = {"resolution": "SK RESOLUTION", "minutes": "MEETING MINUTES", "certificate": "CERTIFICATE OF RECOGNITION"}
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(width / 2, y, title_map.get(template_id, template_id.upper()))
    y -= 30

    # Body — word-wrap long lines so text stays inside the border
    from reportlab.lib.utils import simpleSplit
    _pdf_left = 72
    _pdf_text_w = width - _pdf_left - 72  # usable text width
    c.setFont("Helvetica", 12)
    for line in rendered_text.split("\n"):
        if not line.strip():
            y -= 16
            if y < 72:
                c.showPage()
                _draw_page_border(c)
                y = height - 72
                c.setFont("Helvetica", 12)
            continue
        wrapped = simpleSplit(line, "Helvetica", 12, _pdf_text_w) or ['']
        for wl in wrapped:
            if y < 72:
                c.showPage()
                _draw_page_border(c)
                y = height - 72
                c.setFont("Helvetica", 12)
            c.drawString(_pdf_left, y, wl)
            y -= 16

    c.save()
    buf.seek(0)
    filename = f"{template_id}_{data.get('date', 'document')}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

# ---------------------------------------------------------------------------
# Feature 5: Events Analytics Dashboard  (Phase 2)
# ---------------------------------------------------------------------------
_EVENTS_DB_PATH = os.path.join(os.path.dirname(__file__), "..", "backend", "data", "events.db")

@app.get("/analytics/events")
def analytics_events(type: str = "event", show_staff: bool = False, show_gender: bool = False):
    valid_types = {"event", "monthly", "status", "attendance"}
    if type not in valid_types:
        raise HTTPException(status_code=400, detail=f"Invalid type. Must be one of: {', '.join(valid_types)}")

    db_path = os.path.abspath(_EVENTS_DB_PATH)
    if not os.path.exists(db_path):
        raise HTTPException(status_code=404, detail="Events database not found.")

    import pandas as pd
    import plotly.express as px
    import plotly.graph_objects as go

    conn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True, check_same_thread=False)
    df = pd.read_sql_query("SELECT * FROM events", conn)
    conn.close()

    if df.empty:
        return {"chart": None, "stats": {"total": 0, "upcoming": 0, "completed": 0,
                "total_attendees": 0, "total_male": 0, "total_female": 0,
                "total_staff": 0, "total_budget": 0}}

    completed = df[df["status"] == "completed"]

    stats = {
        "total": int(len(df)),
        "upcoming": int(len(df[df["status"] == "upcoming"])),
        "completed": int(len(completed)),
        "total_attendees": int(completed["attendees"].fillna(0).sum()),
        "total_male": int(completed["male_count"].fillna(0).sum()),
        "total_female": int(completed["female_count"].fillna(0).sum()),
        "total_staff": int(completed["staff_count"].dropna().sum()),
        "total_budget": float(df["budget_allotted"].fillna(0).sum()),
    }

    if type == "event":
        counts = df.groupby("category").size().reset_index(name="count")
        fig = px.pie(counts, names="category", values="count", title="Events by Category")
    elif type == "monthly":
        df["month"] = pd.to_datetime(df["date"], errors="coerce").dt.to_period("M").astype(str)
        counts = df.groupby("month").size().reset_index(name="count")
        fig = px.bar(counts, x="month", y="count", title="Events by Month")
    elif type == "status":
        counts = df.groupby("status").size().reset_index(name="count")
        fig = px.bar(counts, x="status", y="count", title="Events by Status", color="status")
    else:  # attendance
        top = completed.nlargest(10, "attendees")[["title", "attendees", "male_count", "female_count", "staff_count"]].copy()
        top = top.fillna({"attendees": 0, "male_count": 0, "female_count": 0})
        fig = go.Figure()
        fig.add_trace(go.Bar(name="Attendees", x=top["title"], y=top["attendees"]))
        if show_gender:
            fig.add_trace(go.Bar(name="Male", x=top["title"], y=top["male_count"]))
            fig.add_trace(go.Bar(name="Female", x=top["title"], y=top["female_count"]))
        if show_staff:
            staff = top["staff_count"].fillna(0)
            fig.add_trace(go.Bar(name="Staff", x=top["title"], y=staff))
        fig.update_layout(barmode="group", title="Attendance by Event (Top 10)")

    chart_json = fig.to_json()
    return {"chart": chart_json, "stats": stats}

# ---------------------------------------------------------------------------
# Feature 6: Event Document Parser — keyword-based field extraction
# ---------------------------------------------------------------------------
import re as _re

_EVENT_CATEGORIES = ['sports', 'seminar', 'scholarship', 'assembly', 'community',
                     'livelihood', 'general', 'cultural', 'health']

def _extract_fields_from_text(raw_text: str) -> dict:
    """Scan raw_text for event fields using keyword patterns. Returns {extracted, confidence}."""
    lines = raw_text.strip().split('\n')
    text_lower = raw_text.lower()
    extracted = {}
    confidence = {}

    # --- title ---
    m = _re.search(r'(?:event|title|activity|program)\s*[:\-]\s*(.+)', text_lower)
    if m:
        extracted['title'] = m.group(1).strip().title()
        confidence['title'] = 1.0
    elif lines:
        # first non-empty line as fallback
        for ln in lines:
            if ln.strip():
                extracted['title'] = ln.strip()[:120]
                confidence['title'] = 0.4
                break

    # --- date ---
    m = _re.search(r'(?:date)\s*[:\-]\s*(.+)', text_lower)
    if m:
        extracted['date'] = m.group(1).strip()
        confidence['date'] = 1.0
    else:
        # YYYY-MM-DD or Month DD, YYYY or DD/MM/YYYY
        dm = _re.search(r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})', raw_text)
        if not dm:
            dm = _re.search(r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})', raw_text, _re.IGNORECASE)
        if not dm:
            dm = _re.search(r'(\d{1,2}/\d{1,2}/\d{4})', raw_text)
        if dm:
            extracted['date'] = dm.group(1).strip()
            confidence['date'] = 0.7

    # --- time ---
    m = _re.search(r'(?:time)\s*[:\-]\s*(.+)', text_lower)
    if m:
        extracted['time'] = m.group(1).strip()
        confidence['time'] = 1.0
    else:
        tm = _re.search(r'(\d{1,2}:\d{2}\s*(?:AM|PM|am|pm)?)', raw_text)
        if tm:
            extracted['time'] = tm.group(1).strip()
            confidence['time'] = 0.7

    # --- location ---
    m = _re.search(r'(?:venue|location|place|held at)\s*[:\-]\s*(.+)', text_lower)
    if m:
        extracted['location'] = m.group(1).strip().title()
        confidence['location'] = 1.0

    # --- organizer ---
    m = _re.search(r'(?:organized by|organizer|conducted by|facilitated by)\s*[:\-]\s*(.+)', text_lower)
    if m:
        extracted['organizer'] = m.group(1).strip().title()
        confidence['organizer'] = 1.0

    # --- category ---
    for cat in _EVENT_CATEGORIES:
        if cat in text_lower:
            extracted['category'] = cat
            confidence['category'] = 0.9
            break

    # --- attendees ---
    m = _re.search(r'(?:total attendees|number of participants|total participants|attendance)\s*[:\-]\s*(\d+)', text_lower)
    if m:
        extracted['attendees'] = int(m.group(1))
        confidence['attendees'] = 1.0
    else:
        am = _re.search(r'(\d+)\s*(?:attendees|participants)', text_lower)
        if am:
            extracted['attendees'] = int(am.group(1))
            confidence['attendees'] = 0.6

    # --- male_count ---
    m = _re.search(r'(?:male|males|male attendees|male participants)\s*[:\-]\s*(\d+)', text_lower)
    if m:
        extracted['male_count'] = int(m.group(1))
        confidence['male_count'] = 1.0

    # --- female_count ---
    m = _re.search(r'(?:female|females|female attendees|female participants)\s*[:\-]\s*(\d+)', text_lower)
    if m:
        extracted['female_count'] = int(m.group(1))
        confidence['female_count'] = 1.0

    # --- staff_count ---
    m = _re.search(r'(?:staff|facilitators|volunteers)\s*[:\-]\s*(\d+)', text_lower)
    if m:
        extracted['staff_count'] = int(m.group(1))
        confidence['staff_count'] = 1.0

    # --- budget_allotted ---
    m = _re.search(r'(?:budget|allotted|budget allotted)\s*[:\-]\s*(?:PHP|₱|php)?\s*([\d,]+(?:\.\d{1,2})?)', text_lower)
    if m:
        extracted['budget_allotted'] = float(m.group(1).replace(',', ''))
        confidence['budget_allotted'] = 1.0
    else:
        bm = _re.search(r'(?:₱|PHP)\s*([\d,]+(?:\.\d{1,2})?)', raw_text)
        if bm:
            extracted['budget_allotted'] = float(bm.group(1).replace(',', ''))
            confidence['budget_allotted'] = 0.7

    # --- description ---
    for ln in lines:
        stripped = ln.strip()
        if stripped and ':' not in stripped and len(stripped) > 40:
            extracted['description'] = stripped[:500]
            confidence['description'] = 0.5
            break

    # needs_ai: true if any core field has confidence < 0.6
    core_fields = ['title', 'date', 'attendees', 'budget_allotted']
    needs_ai = any(confidence.get(f, 0) < 0.6 for f in core_fields)

    return {"extracted": extracted, "confidence": confidence, "needs_ai": needs_ai}


@app.post("/parse-event-document")
async def parse_event_document(file: UploadFile = File(...)):
    """Extract event fields from an uploaded document via keyword scanning."""
    content_type = file.content_type or ""
    file_bytes = await file.read()
    original_name = file.filename or "upload"
    raw_text = ""

    try:
        # Image files → OCR
        if content_type in IMAGE_MIME_TYPES or content_type.startswith("image/"):
            if not _ocr_available:
                raise HTTPException(status_code=503, detail="Tesseract OCR not available for image parsing.")
            import pytesseract
            from PIL import Image
            img = Image.open(io.BytesIO(file_bytes))
            raw_text = pytesseract.image_to_string(img)

        # PDF → try text extraction, OCR fallback
        elif content_type == "application/pdf":
            if _ocr_available:
                from pdf2image import convert_from_bytes
                import pytesseract
                pages = convert_from_bytes(file_bytes, dpi=300, poppler_path=_poppler_path)
                raw_text = "\n\n".join(pytesseract.image_to_string(p) for p in pages)
            else:
                raw_text = file_bytes.decode("utf-8", errors="replace")

        # DOCX
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or \
             (original_name and original_name.lower().endswith(".docx")):
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_bytes))
            raw_text = "\n".join(p.text for p in doc.paragraphs)

        # Plain text / CSV / Markdown
        else:
            raw_text = file_bytes.decode("utf-8", errors="replace")

    except HTTPException:
        raise
    except Exception as e:
        print(f"[ParseDoc] Text extraction error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text: {str(e)}")

    if not raw_text.strip():
        return {"extracted": {}, "confidence": {}, "raw_text": "", "needs_ai": True}

    result = _extract_fields_from_text(raw_text)
    result["raw_text"] = raw_text
    print(f"[ParseDoc] Extracted {len(result['extracted'])} fields from {original_name} (needs_ai={result['needs_ai']})")
    return result

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
