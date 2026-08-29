"""
document_generator.py
aSK Youth AI — Python Tool 1
─────────────────────────────────────────────────────────────────────────────
Generates formatted DOCX (and optional PDF) SK official documents from
structured field data supplied by the AI via the <TOOL> payload.

The AI never writes document content itself — it extracts field values from
the conversation and sends them here. This service builds the full document,
saving context tokens and ensuring consistent SK formatting.

Supported document types:
  certificate   — Certificate of Participation / Appreciation / Completion
  resolution    — SK Legislative Resolution
  minutes       — Minutes of the Meeting
  letter        — Invitation / Formal Letter
  memo          — Internal Memorandum
  poa           — Program of Activities
  report        — Activity / Accomplishment Report
  proposal      — Project Proposal
  project_brief — ABYIP-format SK Project Brief

API:
  POST /tools/document
  Body: { "type": str, "fields": dict, "language": "filipino"|"english" }
  Returns: { "status": "ok", "filename": str, "download_url": str, "preview": str }

Run standalone:
  python document_generator.py --type certificate --fields '{"recipient_name":"Juan dela Cruz",...}'

Dependencies:
  pip install python-docx flask
"""

import argparse
import json
import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from flask import Flask, request, jsonify

# ── Config ────────────────────────────────────────────────────────────────────
OUTPUT_DIR   = Path(os.environ.get("ASKYOUTH_OUTPUT_DIR", "./generated_docs"))
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

BASE_URL     = os.environ.get("ASKYOUTH_BASE_URL", "http://localhost:5001")
SK_ORG       = "Sangguniang Kabataan"
BARANGAY     = "Barangay Concepcion Dos"
CITY         = "Marikina City"
REPUBLIC     = "Republic of the Philippines"

app = Flask(__name__)

# ── Header builder ────────────────────────────────────────────────────────────
def _add_header(doc: Document) -> None:
    """Adds the standard SK letterhead to a document."""
    for text, bold, size in [
        (REPUBLIC,       False, 11),
        (CITY,           False, 11),
        (BARANGAY,       True,  12),
        (f"OFFICE OF THE {SK_ORG.upper()}", True, 11),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    doc.add_paragraph()  # spacer


def _add_divider(doc: Document) -> None:
    p = doc.add_paragraph("─" * 60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_signature(doc: Document, fields: dict) -> None:
    doc.add_paragraph()
    doc.add_paragraph()
    sig_name  = fields.get("chairperson_name", "[SK Chairperson Name]")
    sig_title = "SK Chairperson"
    sig_org   = f"{SK_ORG}, {BARANGAY}"

    for text, bold in [
        ("________________________________", False),
        (sig_name,  True),
        (sig_title, False),
        (sig_org,   False),
        (CITY,      False),
    ]:
        p = doc.add_paragraph(text)
        if bold:
            for run in p.runs:
                run.bold = True


# ── Document builders ─────────────────────────────────────────────────────────
def build_certificate(fields: dict, language: str) -> Document:
    doc = Document()
    _add_header(doc)

    title = "CERTIFICATE OF PARTICIPATION"
    if "type_of_cert" in fields:
        cert_type = fields["type_of_cert"].upper()
        title = f"CERTIFICATE OF {cert_type}"

    t = doc.add_paragraph(title)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(14)

    doc.add_paragraph()

    name  = fields.get("recipient_name", "[Full Name]")
    event = fields.get("event_name", "[Event Name]")
    date  = fields.get("event_date", "[Date]")
    venue = fields.get("venue", f"[Venue], {BARANGAY}, {CITY}")

    if language == "filipino":
        body = (
            f"Ito ay nagpapatunay na si {name} ay aktibong lumahok sa "
            f"{event} na ginanap noong {date} sa {venue}."
        )
        issued = f"Inilabas ngayong ika-{_today()} sa {BARANGAY}, {CITY}."
    else:
        body = (
            f"This is to certify that {name} has actively participated in "
            f"{event} held on {date} at {venue}."
        )
        issued = f"Issued this {_today(long=True)} at {BARANGAY}, {CITY}."

    doc.add_paragraph(body)
    doc.add_paragraph()
    doc.add_paragraph(issued)

    _add_signature(doc, fields)
    return doc


def build_resolution(fields: dict, language: str) -> Document:
    doc = Document()
    _add_header(doc)

    number = fields.get("resolution_number", "[No.]")
    series = fields.get("series", datetime.now().year)
    title  = fields.get("title", "[Resolution Title]")
    date   = fields.get("meeting_date", _today(long=True))

    heading = f"RESOLUTION NO. {number}, SERIES OF {series}"
    p = doc.add_paragraph(heading)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    sub = doc.add_paragraph(f'"{title}"')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(f"Approved during the SK Session held on {date}.")
    doc.add_paragraph()

    whereas_clauses = fields.get("whereas_clauses", [
        "the Sangguniang Kabataan is mandated to serve the youth of the barangay;",
        "this resolution supports youth development in accordance with RA 10742;"
    ])
    for clause in whereas_clauses:
        w = doc.add_paragraph()
        w.add_run("WHEREAS, ").bold = True
        w.add_run(clause)

    doc.add_paragraph()
    resolve = doc.add_paragraph()
    resolve.add_run("NOW, THEREFORE BE IT RESOLVED, ").bold = True
    resolve.add_run(fields.get(
        "resolved_clause",
        "as it is hereby resolved to approve the foregoing matter."
    ))

    doc.add_paragraph()
    doc.add_paragraph("APPROVED UNANIMOUSLY.")
    _add_signature(doc, fields)
    return doc


def build_minutes(fields: dict, language: str) -> Document:
    doc = Document()
    _add_header(doc)

    title = "MINUTES OF THE MEETING"
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph(f"Date    : {fields.get('meeting_date', fields.get('date', '[Date]'))}")
    doc.add_paragraph(f"Time    : {fields.get('meeting_time', fields.get('time', '[Time]'))}")
    doc.add_paragraph(f"Venue   : {fields.get('venue', '[Venue]')}")
    doc.add_paragraph(f"Presided by : {fields.get('presided_by', fields.get('presiding_officer', '[SK Chairperson]'))}")

    doc.add_paragraph()
    p = doc.add_paragraph("I. ATTENDEES"); p.runs[0].bold = True
    attendees = fields.get("attendees", ["[Officer 1]", "[Officer 2]"])
    if isinstance(attendees, str):
        attendees = [a.strip() for a in attendees.split('\n') if a.strip()]
    for a in attendees:
        doc.add_paragraph(f"    • {a}")

    doc.add_paragraph()
    p = doc.add_paragraph("II. AGENDA AND DISCUSSIONS"); p.runs[0].bold = True
    
    agenda_items = fields.get("agenda_items", [])
    if isinstance(agenda_items, str):
        agenda_items = [a.strip() for a in agenda_items.split('\n') if a.strip()]

    if not agenda_items:
        agenda_items = ["[Agenda item 1]"]

    for i, item in enumerate(agenda_items, 1):
        if isinstance(item, dict):
            # Rich format: {item, time_start, time_end, discussion}
            topic = item.get("item", item.get("topic", item.get("agenda", "[Topic]")))
            t_start = item.get("time_start", "")
            t_end = item.get("time_end", "")
            time_range = f"  ({t_start} – {t_end})" if t_start and t_end else (f"  ({t_start})" if t_start else "")
            
            # Agenda title line — bold topic + time
            p = doc.add_paragraph()
            run_num = p.add_run(f"    {i}. ")
            run_title = p.add_run(topic)
            run_title.bold = True
            if time_range:
                run_time = p.add_run(time_range)
                run_time.bold = False
                run_time.font.size = Pt(10)

            # Discussion narrative
            discussion = item.get("discussion", "")
            if discussion:
                dp = doc.add_paragraph(f"       {discussion}")
                dp.paragraph_format.left_indent = Inches(0.5)
                dp.paragraph_format.space_after = Pt(6)
        else:
            # Simple string fallback
            doc.add_paragraph(f"    {i}. {item}")

    # Standalone proceedings if AI used the old flat field
    proceedings = fields.get("proceedings", "")
    if proceedings and not any(isinstance(x, dict) for x in agenda_items):
        doc.add_paragraph()
        p = doc.add_paragraph("III. PROCEEDINGS / DISCUSSION"); p.runs[0].bold = True
        if isinstance(proceedings, list):
            for d in proceedings:
                doc.add_paragraph(f"    {d}")
        else:
            doc.add_paragraph(f"    {proceedings}")

    doc.add_paragraph()
    p = doc.add_paragraph("IV. ACTION ITEMS"); p.runs[0].bold = True
    action_items = fields.get("action_items", fields.get("decisions", ["[Action 1]"]))
    if isinstance(action_items, str):
        action_items = [a.strip() for a in action_items.split('\n') if a.strip()]
    for a in action_items:
        doc.add_paragraph(f"    • {a}")

    doc.add_paragraph()
    p = doc.add_paragraph("V. ADJOURNMENT"); p.runs[0].bold = True
    adj_time = fields.get("adjournment_time", fields.get("adjournment", "[Time]"))
    doc.add_paragraph(f"    The meeting was adjourned at {adj_time}.")

    doc.add_paragraph()
    doc.add_paragraph()
    
    prepared_by = fields.get("prepared_by", "[SK Secretary Name]")
    noted_by = fields.get("noted_by", fields.get("chairperson_name", "[SK Chairperson Name]"))
    
    doc.add_paragraph("Prepared by:")
    doc.add_paragraph()
    p = doc.add_paragraph(prepared_by); p.runs[0].bold = True
    doc.add_paragraph("SK Secretary")
    
    doc.add_paragraph()
    doc.add_paragraph("Noted by:")
    doc.add_paragraph()
    p = doc.add_paragraph(noted_by); p.runs[0].bold = True
    doc.add_paragraph("SK Chairperson")

    return doc


def build_letter(fields: dict, language: str) -> Document:
    doc = Document()
    _add_header(doc)

    doc.add_paragraph(_today(long=True))
    doc.add_paragraph()
    doc.add_paragraph(fields.get("recipient_name", "[Recipient Name]"))
    doc.add_paragraph(fields.get("recipient_address", "[Address]"))
    doc.add_paragraph()

    salutation = fields.get("salutation", "Dear Sir/Ma'am,")
    doc.add_paragraph(salutation)
    doc.add_paragraph()
    doc.add_paragraph(fields.get("body", "[Letter body here.]"))
    doc.add_paragraph()
    doc.add_paragraph("Respectfully yours,")
    _add_signature(doc, fields)
    return doc


def build_memo(fields: dict, language: str) -> Document:
    doc = Document()
    _add_header(doc)

    t = doc.add_paragraph("MEMORANDUM")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph(f"TO      : {fields.get('to', '[Recipient]')}")
    doc.add_paragraph(f"FROM    : {fields.get('from_name', '[Sender]')}")
    doc.add_paragraph(f"DATE    : {_today(long=True)}")
    doc.add_paragraph(f"SUBJECT : {fields.get('subject', '[Subject]')}")
    doc.add_paragraph()
    doc.add_paragraph(fields.get("body", "[Memo body here.]"))
    _add_signature(doc, fields)
    return doc


def build_poa(fields: dict, language: str) -> Document:
    doc = Document()
    _add_header(doc)

    t = doc.add_paragraph("PROGRAM OF ACTIVITIES")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True

    doc.add_paragraph(f"Event  : {fields.get('event_name', '[Event]')}")
    doc.add_paragraph(f"Date   : {fields.get('event_date', '[Date]')}")
    doc.add_paragraph(f"Venue  : {fields.get('venue', '[Venue]')}")
    doc.add_paragraph()

    activities = fields.get("activities", [
        {"time": "8:00 AM",  "activity": "Registration"},
        {"time": "9:00 AM",  "activity": "Opening Program"},
        {"time": "12:00 NN", "activity": "Lunch Break"},
        {"time": "1:00 PM",  "activity": "Main Activity"},
        {"time": "5:00 PM",  "activity": "Closing Remarks"},
    ])

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "TIME"
    hdr[1].text = "ACTIVITY"
    for act in activities:
        row = table.add_row().cells
        row[0].text = act.get("time", "")
        row[1].text = act.get("activity", "")

    _add_signature(doc, fields)
    return doc


def build_report(fields: dict, language: str) -> Document:
    doc = Document()
    _add_header(doc)

    report_type = fields.get("report_type", "ACTIVITY REPORT").upper()
    t = doc.add_paragraph(report_type)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph(f"Event          : {fields.get('event_name', '[Event]')}")
    doc.add_paragraph(f"Date           : {fields.get('event_date', '[Date]')}")
    doc.add_paragraph(f"Venue          : {fields.get('venue', '[Venue]')}")
    doc.add_paragraph(f"Participants   : {fields.get('participant_count', '[#]')}")
    doc.add_paragraph(f"Budget Used    : {fields.get('budget_used', '[Amount]')}")
    doc.add_paragraph()

    sections = [
        ("I. OBJECTIVES",   "objectives"),
        ("II. NARRATIVE",   "narrative"),
        ("III. OUTCOMES",   "outcomes"),
        ("IV. RECOMMENDATIONS", "recommendations"),
    ]
    for heading, key in sections:
        p = doc.add_paragraph(heading); p.runs[0].bold = True
        doc.add_paragraph(fields.get(key, f"[{heading.split('. ',1)[1].title()}]"))
        doc.add_paragraph()

    _add_signature(doc, fields)
    return doc


def _set_run_font(run, size=Pt(12), bold=False, underline=False, color=None):
    """Force Times New Roman + size on a single run. python-docx does NOT cascade
    the Normal style font into table cells or headers, so every run needs this."""
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.bold = bold
    run.underline = underline
    # East-Asian fallback — required for Word to actually render Times New Roman
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if color:
        run.font.color.rgb = color


def build_project_brief(fields: dict, language: str) -> Document:
    """Generates a Project Brief in the official SK ABYIP barangay format."""
    doc = Document()

    # 1. Global Typography — set Normal style as baseline
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # 2. Official Letterhead (in the Document Header, NOT the body)
    section = doc.sections[0]
    section.header_distance = Cm(1.27)
    header = section.header
    header.is_linked_to_previous = False

    # Clear default empty paragraph
    for p in header.paragraphs:
        p.clear()

    # Line 1: Republic of the Philippines (11pt, black)
    p1 = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run('Republic of the Philippines')
    _set_run_font(r1, size=Pt(11))

    # Line 2: SANGGUNIANG KABATAAN (16pt, red)
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    r2 = p2.add_run('SANGGUNIANG KABATAAN')
    _set_run_font(r2, size=Pt(16), color=RGBColor(218, 41, 28))

    # Line 3: BARANGAY CONCEPCION DOS (16pt, red)
    p3 = header.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.space_before = Pt(0)
    r3 = p3.add_run('BARANGAY CONCEPCION DOS')
    _set_run_font(r3, size=Pt(16), color=RGBColor(218, 41, 28))

    # Line 4: City of Marikina... (11pt, black)
    p4 = header.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(0)
    r4 = p4.add_run('City of Marikina, Metro Manila, Philippines')
    _set_run_font(r4, size=Pt(11))

    # 3. Document Title in body
    doc.add_paragraph()  # spacer
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = t.add_run('PROJECT BRIEF')
    _set_run_font(t_run, size=Pt(12), bold=True)
    doc.add_paragraph()  # spacer

    # 4. Two-Column Project Table with full grid borders
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'

    rows_data = [
        ('NAME OF PROJECT/ACTIVITY',       fields.get('project_name',    '[Project Name]'),            True),
        ('LOCATION OF PROJECT',            fields.get('location',        'Barangay Concepcion Dos'),   False),
        ('TARGET DATE OF IMPLEMENTATION',  fields.get('target_date',     '[Target Date]'),             False),
        ('BACKGROUND/RATIONALE',           fields.get('background',      '[Background Narrative]'),    False),
        ('OBJECTIVE',                      fields.get('objective',       '[Objective Narrative]'),     False),
        ('TARGET PHYSICAL OUTPUT',         fields.get('physical_output', '[Target Physical Output]'),  False),
        ('TARGET BENEFICIARIES',           fields.get('beneficiaries',   '[Target Beneficiaries]'),    False),
        ('BUDGET:',                        f"Php {fields.get('budget',   '[Amount]')}",               False),
    ]

    for i, (label, val, bold_val) in enumerate(rows_data):
        # Left column — label
        cell_left = table.rows[i].cells[0]
        p_left = cell_left.paragraphs[0]
        p_left.clear()
        r_left = p_left.add_run(label)
        _set_run_font(r_left, size=Pt(12), bold=True)

        # Right column — value
        cell_right = table.rows[i].cells[1]
        p_right = cell_right.paragraphs[0]
        p_right.clear()
        r_right = p_right.add_run(val)
        _set_run_font(r_right, size=Pt(12), bold=bold_val)

    doc.add_paragraph()  # spacer
    doc.add_paragraph()  # spacer

    # 5. Official Signature Section
    p_prep = doc.add_paragraph()
    r_prep = p_prep.add_run('Prepared by:')
    _set_run_font(r_prep, size=Pt(12))
    doc.add_paragraph()
    doc.add_paragraph()

    # Lead Official (centered)
    p_lead = doc.add_paragraph()
    p_lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lead = p_lead.add_run(fields.get('prepared_by_name', 'LOUIE MARI LAMPA').upper())
    _set_run_font(r_lead, size=Pt(12), bold=True, underline=True)

    p_lead_t1 = doc.add_paragraph()
    p_lead_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t1 = p_lead_t1.add_run('Sangguniang Kabataan Chairperson')
    _set_run_font(r_t1, size=Pt(12))

    p_lead_t2 = doc.add_paragraph()
    p_lead_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t2 = p_lead_t2.add_run('Chairperson, Committee on Youth and Sports Development')
    _set_run_font(r_t2, size=Pt(12))

    doc.add_paragraph()
    doc.add_paragraph()

    # SK Council Member Signature Grid (6 in 2x3, 7th centered)
    members = [
        'JOSHUA E. AROCENA',         'JEAN B. CABILING',
        'MA. LOUELLA B. MENDOZA',    'REINNIER B. YARZA',
        'JULIA BEATRICE F. OLAYVAR', 'CLYDE HARRY S. GOMEZ',
        'TRICIA YZABEL T. SULIT',
    ]

    sig_table = doc.add_table(rows=3, cols=2)
    # Remove borders from signature table
    for row in sig_table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'none')
                el.set(qn('w:sz'), '0')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), 'auto')
                tcBorders.append(el)
            tcPr.append(tcBorders)

    for idx in range(6):
        r_idx = idx // 2
        c_idx = idx % 2
        cell_p = sig_table.rows[r_idx].cells[c_idx].paragraphs[0]
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r_name = cell_p.add_run(members[idx])
        _set_run_font(r_name, size=Pt(12), bold=True, underline=True)

        r_sub = cell_p.add_run('\nSangguniang Kabataan Member\n\n')
        _set_run_font(r_sub, size=Pt(12))

    # 7th member centered
    p_last = doc.add_paragraph()
    p_last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_last_name = p_last.add_run(members[6])
    _set_run_font(r_last_name, size=Pt(12), bold=True, underline=True)
    r_last_sub = p_last.add_run('\nSangguniang Kabataan Member')
    _set_run_font(r_last_sub, size=Pt(12))

    return doc


def build_proposal(fields: dict, language: str) -> Document:
    doc = Document()
    _add_header(doc)

    t = doc.add_paragraph("PROJECT PROPOSAL")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True

    doc.add_paragraph(f"Project Title  : {fields.get('project_title', '[Title]')}")
    doc.add_paragraph(f"Proponent      : {SK_ORG}, {BARANGAY}")
    doc.add_paragraph(f"Target Date    : {fields.get('target_date', '[Date]')}")
    doc.add_paragraph(f"Target Beneficiaries : {fields.get('beneficiaries', '[Beneficiaries]')}")
    doc.add_paragraph(f"Estimated Budget : {fields.get('budget', '[Amount]')}")
    doc.add_paragraph()

    sections = [
        ("I. RATIONALE",    "rationale"),
        ("II. OBJECTIVES",  "objectives"),
        ("III. ACTIVITIES", "activities_narrative"),
        ("IV. TIMELINE",    "timeline"),
        ("V. BUDGET BREAKDOWN", "budget_breakdown"),
    ]
    for heading, key in sections:
        p = doc.add_paragraph(heading); p.runs[0].bold = True
        doc.add_paragraph(fields.get(key, f"[{heading.split('. ',1)[1].title()}]"))
        doc.add_paragraph()

    _add_signature(doc, fields)
    return doc


# ── Dispatch table ────────────────────────────────────────────────────────────
BUILDERS = {
    "certificate":   build_certificate,
    "resolution":    build_resolution,
    "minutes":       build_minutes,
    "letter":        build_letter,
    "memo":          build_memo,
    "poa":           build_poa,
    "report":        build_report,
    "proposal":      build_proposal,
    "project_brief": build_project_brief,
}


# ── Core generate function ────────────────────────────────────────────────────
def generate_document(doc_type: str, fields: dict, language: str = "english") -> dict:
    """
    Generate an SK document and save to OUTPUT_DIR.
    Returns a dict with filename, download_url, and a short preview string.
    """
    doc_type = doc_type.lower().strip()
    language = language.lower().strip()

    if doc_type not in BUILDERS:
        return {"status": "error", "message": f"Unknown document type: {doc_type}"}

    builder  = BUILDERS[doc_type]
    doc      = builder(fields, language)
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"SK_{doc_type}_{ts}.docx"
    filepath = OUTPUT_DIR / filename

    doc.save(str(filepath))

    preview = _build_preview(doc_type, fields)

    return {
        "status":       "ok",
        "filename":     filename,
        "download_url": f"{BASE_URL}/downloads/{filename}",
        "preview":      preview,
    }


def _build_preview(doc_type: str, fields: dict) -> str:
    """Returns a short human-readable summary of what was generated."""
    name  = fields.get("recipient_name") or fields.get("project_title") or ""
    event = fields.get("event_name", "")
    date  = fields.get("event_date") or fields.get("meeting_date") or _today()
    snippets = {
        "certificate": f"Certificate for {name} — {event} ({date})",
        "resolution":  f"Resolution No. {fields.get('resolution_number','?')}, Series {fields.get('series', datetime.now().year)}",
        "minutes":     f"Minutes of the Meeting — {date}",
        "letter":      f"Letter to {fields.get('recipient_name','?')} — {date}",
        "memo":        f"Memo: {fields.get('subject','?')} — {date}",
        "poa":         f"Program of Activities — {event} ({date})",
        "report":      f"{fields.get('report_type','Activity Report')} — {event} ({date})",
        "proposal":      f"Project Proposal: {fields.get('project_title','?')}",
        "project_brief": f"Project Brief: {fields.get('project_name','?')} — {fields.get('target_date', '')}",
    }
    return snippets.get(doc_type, f"{doc_type.title()} document generated.")


def _today(long: bool = False) -> str:
    now = datetime.now()
    if long:
        day = now.day
        suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
        return now.strftime(f"{day}{suffix} day of %B, %Y")
    return now.strftime("%B %d, %Y")


# ── Flask API ──────────────────────────────────────────────────────────────────
@app.route("/tools/document", methods=["POST"])
def api_generate():
    data     = request.get_json(force=True)
    doc_type = data.get("type", "")
    fields   = data.get("fields", {})
    language = data.get("language", "english")

    if not doc_type:
        return jsonify({"status": "error", "message": "Missing 'type' field."}), 400

    result = generate_document(doc_type, fields, language)
    status_code = 200 if result["status"] == "ok" else 400
    return jsonify(result), status_code


@app.route("/downloads/<filename>")
def serve_file(filename):
    from flask import send_from_directory
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)


# ── CLI entry point ────────────────────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="aSK Youth Document Generator")
    parser.add_argument("--serve",    action="store_true", help="Run as Flask API server")
    parser.add_argument("--port",     type=int, default=5001)
    parser.add_argument("--type",     type=str, default="certificate")
    parser.add_argument("--fields",   type=str, default="{}")
    parser.add_argument("--language", type=str, default="english")
    args = parser.parse_args()

    if args.serve:
        print(f"[DocGen] Serving on port {args.port}")
        app.run(host="0.0.0.0", port=args.port, debug=False)
    else:
        fields = json.loads(args.fields)
        result = generate_document(args.type, fields, args.language)
        print(json.dumps(result, indent=2))
